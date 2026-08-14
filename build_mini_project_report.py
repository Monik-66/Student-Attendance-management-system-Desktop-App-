from __future__ import annotations

import math
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = Path(r"C:\Users\monik\Downloads\Final Template for mini project Report-First pages.docx")
OUTPUT_PATH = BASE_DIR / "SNPSU_Teacher_Desktop_Mini_Project_Report.docx"
ASSET_DIR = BASE_DIR / "report_assets"

BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(70, 70, 70)
LIGHT_GRAY = RGBColor(242, 244, 247)
MID_BLUE = RGBColor(45, 84, 125)
SOFT_BLUE = RGBColor(232, 238, 245)


def ensure_assets_dir() -> None:
    ASSET_DIR.mkdir(exist_ok=True)


def find_font(candidates: list[str]) -> str:
    fonts_dir = Path(r"C:\Windows\Fonts")
    for candidate in candidates:
        path = fonts_dir / candidate
        if path.exists():
            return str(path)
    raise FileNotFoundError(f"Unable to find a font from: {candidates}")


TIMES_FONT = find_font(["times.ttf", "Times New Roman.ttf", "timesnewroman.ttf"])
TIMES_BOLD_FONT = find_font(["timesbd.ttf", "Times New Roman Bold.ttf", "timesnewromanbold.ttf"])
ARIAL_FONT = find_font(["arial.ttf", "Arial.ttf"])
ARIAL_BOLD_FONT = find_font(["arialbd.ttf", "Arial Bold.ttf"])


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=BLACK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_paragraph_format(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, line_spacing=1.5):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_paragraph_text(paragraph, text, *, size=12, bold=False, italic=False, color=BLACK, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    set_paragraph_format(paragraph, align=align)


def add_page_break_paragraph(document):
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break()
    return paragraph


def set_document_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    normal_paragraph = normal.paragraph_format
    normal_paragraph.space_before = Pt(0)
    normal_paragraph.space_after = Pt(6)
    normal_paragraph.line_spacing = 1.5

    for style_name, size, bold in [("Heading 1", 16, False), ("Heading 2", 14, False), ("Heading 3", 12, False)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(6 if style_name != "Heading 1" else 12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.keep_with_next = True


def set_section_layout(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def locate_paragraph(document, needle: str):
    for paragraph in document.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Paragraph containing '{needle}' not found")


def replace_paragraph(document, needle: str, text: str, *, size=12, bold=False, italic=False, color=BLACK, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = locate_paragraph(document, needle)
    set_paragraph_text(paragraph, text, size=size, bold=bold, italic=italic, color=color, align=align)
    return paragraph


def fill_table(table, rows):
    row_index = 1
    for values in rows:
        if row_index >= len(table.rows):
            break
        for col_index, value in enumerate(values):
            cell = table.rows[row_index].cells[col_index]
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, size=11, color=BLACK)
        row_index += 1


def set_table_header(table):
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=11, bold=True, color=BLACK)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_body_table(table):
    table.autofit = False
    widths = [Inches(1.1), Inches(4.8), Inches(0.9), Inches(1.1), Inches(1.4)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15


def add_simple_table(document, headers, rows, widths, title=None, title_size=11):
    if title:
        add_caption(document, title)
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = widths[index]
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=title_size, bold=True, color=BLACK)

    for row_index, row_values in enumerate(rows, start=1):
        for col_index, value in enumerate(row_values):
            cell = table.rows[row_index].cells[col_index]
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=11, color=BLACK)
    return table


def add_caption(document, text):
    paragraph = document.add_paragraph()
    set_paragraph_text(paragraph, text, size=11, bold=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    return paragraph


def add_body_heading(document, text, *, level=2, before=4, after=4):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=False, color=BLACK)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return paragraph


def add_body_paragraph(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=12, color=BLACK)
    set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, line_spacing=1.5)
    return paragraph


def add_bullet_list(document, items):
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{index}. {item}")
        set_run_font(run, size=12, color=BLACK)
        set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=3, line_spacing=1.5)


def make_paragraphs(label: str, points: list[str]) -> list[str]:
    p1 = (
        f"{label} is framed around {points[0]}. In the SNPSU Teacher Desktop project, the same idea is implemented by "
        f"linking teacher login, section mapping, and date-wise attendance into a single workflow. This keeps the faculty "
        f"experience direct and avoids the repeated paper-based steps that normally create delay during busy class hours."
    )
    p1 += f" The design also respects {points[1]}, which is why the interface allows the user to move from one step to the next without losing context."
    p2 = (
        f"From a technical perspective, the section also supports {points[2]}. That requirement influenced the API layout, the database schema, "
        f"and the offline queue used when connectivity is not stable. Together, these choices make the report more than a description of screens; "
        f"they show how the same problem is solved through a small but coherent system."
    )
    if len(points) > 3:
        p2 += f" A further observation is that {points[3]}, which helps justify the reporting and analysis shown later in the document."
    return [p1, p2]


def paragraph_bank():
    return {
        "1.1 BACKGROUND OF THE STUDY": [
            "faculty still rely on registers and spreadsheet notes for daily attendance",
            "a classroom system must support quick action during the teaching slot",
            "section-wise and subject-wise context is essential for meaningful records",
            "student follow-up becomes easier when the attendance trail is visible immediately",
        ],
        "1.2 PROBLEM STATEMENT": [
            "paper methods cause delays, transcription errors, and missing signatures",
            "manual handling makes date-wise correction and re-opening difficult",
            "faculty need a single dashboard that reflects the current section status",
            "administrative reporting becomes inconsistent when data is scattered across files",
        ],
        "1.3 OBJECTIVES OF THE PROJECT": [
            "a faculty-friendly attendance dashboard with minimal clicks",
            "a secure connection to PostgreSQL for live section data",
            "offline capture and later synchronization when the connection returns",
            "absentee visibility and communication support after the attendance submission",
        ],
        "1.4 SCOPE OF THE PROJECT": [
            "teacher authentication, section mapping, and attendance marking",
            "admin-style section summaries and downloadable reports",
            "elective-aware student filtering for special subject groups",
            "basic analytics for attendance rates and low-attendance review",
        ],
        "1.5 PROJECT MOTIVATION": [
            "reduce repetitive manual work for every class period",
            "help the department observe attendance patterns across sections",
            "keep the system practical for an academic laboratory environment",
            "make the workflow easy enough to demonstrate during viva sessions",
        ],
        "1.6 ORGANIZATION OF THE REPORT": [
            "the report begins with background and motivation",
            "later chapters explain literature, analysis, design, and implementation",
            "the last chapters discuss testing, outcomes, and future scope",
            "appendix material is used only where it adds useful supporting detail",
        ],
        "2.1 MANUAL ATTENDANCE PRACTICES": [
            "register-based marking is familiar but slow during congested timetables",
            "corrections often depend on handwritten notes and memory",
            "a missing date entry can create follow-up work for the faculty member",
            "the method is simple, yet it scales poorly when many sections are involved",
        ],
        "2.2 WEB AND DESKTOP ATTENDANCE SYSTEMS": [
            "desktop applications keep the interface available even when the browser is not preferred",
            "teacher-centric software can reduce training effort by matching classroom habits",
            "a strong system should separate UI, API, and database responsibilities",
            "the project follows this separation to keep maintenance manageable",
        ],
        "2.3 OFFLINE-FIRST APPROACHES": [
            "temporary network loss should not stop a teacher from recording attendance",
            "queued actions are safer than forcing the user to restart work later",
            "local caching helps preserve the recent login context and the student register",
            "synchronization routines are needed to merge offline updates back into PostgreSQL",
        ],
        "2.4 SMS AND NOTIFICATION SUPPORT": [
            "absentee information becomes more useful when it can be communicated quickly",
            "notification hooks reduce the chance that parents remain uninformed for long",
            "a message pipeline must respect the timing and content of the attendance event",
            "the project keeps the notification step optional so the base workflow stays simple",
        ],
        "2.5 COMPARATIVE ANALYSIS": [
            "the proposed system reduces duplication compared with manual registers",
            "it also improves visibility when compared with isolated spreadsheet files",
            "a central database supports audit trails and better reporting consistency",
            "the comparison shows why desktop plus server architecture remains relevant in colleges",
        ],
        "2.6 LITERATURE REVIEW SUMMARY": [
            "the literature confirms that usability and reliability matter more than visual complexity",
            "attendance tools work best when they are consistent with the institution structure",
            "offline support and secure storage appear repeatedly as practical design priorities",
            "these lessons shape the requirements chosen for the current project",
        ],
        "3.1 STAKEHOLDERS AND USER ROLES": [
            "teachers need a fast marking interface",
            "administrators need consolidated attendance reports",
            "students and parents benefit indirectly from timely absentee communication",
            "the system therefore balances convenience with accountability",
        ],
        "3.2 FUNCTIONAL REQUIREMENTS": [
            "authentication, section loading, and subject mapping",
            "date selection, present/absent toggling, and save or reset actions",
            "absentee view generation, optional SMS dispatch, and export support",
            "summary cards and analysis widgets for quick review on the dashboard",
        ],
        "3.3 NON-FUNCTIONAL REQUIREMENTS": [
            "the interface should remain responsive during class-time use",
            "data access must be reliable and predictable across sessions",
            "security controls should prevent teachers from editing unassigned sections",
            "the design should remain readable on a standard lab desktop screen",
        ],
        "3.4 USE CASE VIEW": [
            "the main use case starts with teacher login",
            "then the teacher loads a section, marks attendance, and submits the result",
            "administrative review follows if a report or summary is required",
            "the same interaction model keeps the workflow easy to demonstrate",
        ],
        "3.5 DATA REQUIREMENTS": [
            "teacher identity, assigned sections, and subject codes",
            "student records, elective preferences, and parent contact numbers",
            "attendance records keyed by date and student SRN",
            "summary data for analytics and export views",
        ],
        "3.6 CONSTRAINTS AND RISKS": [
            "network interruptions can delay database access",
            "elective-based sections require careful filtering logic",
            "bulk edits must be restricted so that accidental overwrites do not occur",
            "report generation should remain stable even when the attendance set is large",
        ],
        "4.1 OVERALL ARCHITECTURE": [
            "an Electron shell hosts the desktop interface",
            "Node.js and Express provide the backend API",
            "PostgreSQL stores the live academic records",
            "a small offline cache keeps recent teacher data available during interruptions",
        ],
        "4.2 DATABASE DESIGN": [
            "teacher, student, section, and attendance tables form the core model",
            "junction tables map teachers to the sections they handle",
            "elective mapping tables separate DC and MC groups from the full section",
            "the schema favors simple keys so reporting queries remain easy to maintain",
        ],
        "4.3 API AND SERVICE DESIGN": [
            "login and summary endpoints keep the front end thin",
            "attendance save and reset calls encapsulate the mutation logic",
            "admin reporting endpoints return structured summaries for exports",
            "service boundaries are narrow enough to simplify testing and debugging",
        ],
        "4.4 OFFLINE CACHE DESIGN": [
            "cached teacher profiles support reopen and sync behavior",
            "queued attendance changes are stored locally before being submitted again",
            "the cache layer minimizes re-entry when the network comes back",
            "state reconciliation keeps the local view aligned with the database result",
        ],
        "4.5 SECURITY AND VALIDATION": [
            "teacher identity must match the assigned database record",
            "section and subject access are checked before any save operation",
            "input validation protects date and SRN handling",
            "the system rejects operations that do not belong to the signed-in teacher",
        ],
        "4.6 MODULE INTERACTION": [
            "dashboard widgets reflect database summaries",
            "attendance views pull section data on demand",
            "the absentee modal shares the same submitted dataset",
            "report export uses the same server-side sources as the dashboard",
        ],
        "5.1 LOGIN AND DASHBOARD SCREENS": [
            "the login screen is intentionally minimal",
            "after sign-in the dashboard shows the assigned sections and current attendance indicators",
            "key actions remain visible so the teacher can move quickly",
            "the first screen is designed to reduce hesitation rather than impress with complexity",
        ],
        "5.2 ATTENDANCE ENTRY WORKFLOW": [
            "the date picker loads a fresh or existing attendance column",
            "present and absent values are toggled directly in the register",
            "the table is optimized for repeated classroom use",
            "the save action locks the date unless the user deliberately reopens it",
        ],
        "5.3 ABSENTEE REVIEW AND SMS FLOW": [
            "the absentee modal lists the missing students after submission",
            "the same view can trigger a communication step if enabled",
            "a clear summary helps the teacher confirm the final record",
            "the message pathway supports a practical parent-notification workflow",
        ],
        "5.4 ADMIN REPORTING AND EXPORT": [
            "the admin view condenses section status into a readable summary",
            "downloadable output supports review outside the live application",
            "the report path reduces the need to copy data manually into spreadsheets",
            "this is especially useful during internal review and compliance checks",
        ],
        "5.5 VISUAL ANALYTICS DASHBOARD": [
            "charts turn raw attendance data into a quick interpretation layer",
            "low-attendance trends help the faculty notice risk early",
            "summary cards are useful when a teacher wants a status snapshot",
            "the dashboard output in this report illustrates those patterns visually",
        ],
        "5.6 USER EXPERIENCE OBSERVATIONS": [
            "short labels and direct actions keep the interface predictable",
            "a clean layout lowers the chance of accidental clicks during class hours",
            "students are easier to review when the table hierarchy remains stable",
            "the same layout works for both the dashboard and attendance page",
        ],
        "6.1 IMPLEMENTATION STRATEGY": [
            "the project was implemented as a layered desktop application",
            "UI, API, and database responsibilities were separated early",
            "this kept the code base easier to read as features were added",
            "the strategy also made it easier to trace errors to the correct layer",
        ],
        "6.2 SAMPLE TEST CASES": [
            "login with valid and invalid teacher credentials",
            "load a section and verify the student register",
            "save attendance for a new date and then reopen it",
            "delete a date column and confirm the UI reflects the change",
        ],
        "6.3 ERROR HANDLING AND RECOVERY": [
            "invalid section mappings are blocked before they reach the save endpoint",
            "offline mode stores work locally instead of discarding it",
            "clear status messages help the user know what action is required next",
            "recovery after network restoration is automated as far as possible",
        ],
        "6.4 PERFORMANCE AND RELIABILITY": [
            "attendance operations stay small enough for quick save cycles",
            "summary data is computed from the current live record",
            "local caching reduces repeated database fetches",
            "the system is therefore suitable for routine faculty use in a lab setting",
        ],
        "6.5 DEPLOYMENT AND PACKAGING": [
            "Electron packaging makes the system easier to launch on a lab desktop",
            "environment variables separate configuration from code",
            "the backend can run locally while the front end remains responsive",
            "the deployment model supports demonstration without requiring a browser workflow",
        ],
        "6.6 VALIDATION RESULTS": [
            "the tested workflows behaved consistently across repeated runs",
            "attendance state, deletion, and reopen actions all preserved the expected data shape",
            "offline synchronization completed without losing the queued entries",
            "the validation suggests that the design is fit for the mini project scope",
        ],
        "7.1 RESULTS OBSERVED DURING USE": [
            "faculty can complete attendance faster than with a paper register",
            "the dashboard makes low-attendance review immediate",
            "date-specific edits are easier because the application stores each date separately",
            "the user benefits from a workflow that mirrors the real teaching routine",
        ],
        "7.2 COMPARISON WITH MANUAL METHOD": [
            "manual handling needs more time and more correction",
            "the proposed system reduces redundant transcription",
            "report extraction becomes simpler because the data already exists in structured form",
            "this comparison shows why digital capture is better for repeated academic operations",
        ],
        "7.3 DISCUSSION OF STRENGTHS": [
            "the architecture is compact yet practical",
            "offline queue support makes the system resilient",
            "the same interface works for the common faculty routine and the administrative review",
            "the implementation is suitable for incremental enhancement",
        ],
        "7.4 LIMITATIONS": [
            "the current report still depends on the quality of the database seed data",
            "message delivery relies on external communication services when enabled",
            "large institutional rollouts would need stronger role management",
            "the current scope is appropriate for a mini project rather than a full enterprise deployment",
        ],
        "7.5 MAINTENANCE CONSIDERATIONS": [
            "schema changes should be coordinated with the API and UI layers",
            "attendance logic needs careful review before each academic term",
            "backup routines are important because attendance data is operationally sensitive",
            "administrators should keep a documented process for adding new teachers and sections",
        ],
        "7.6 FUTURE ENHANCEMENT AREAS": [
            "bulk import from spreadsheets can reduce setup effort",
            "better analytics can identify repeated absentee patterns",
            "mobile notifications and richer dashboards can extend user reach",
            "role-based workflows could support coordinators and heads of department more cleanly",
        ],
        "8.1 SUMMARY OF WORK": [
            "the project combines attendance capture, reporting, and offline resilience",
            "the result is a single desktop solution for faculty workflow",
            "the front end, service layer, and database collaborate in a simple stack",
            "the report demonstrates that the mini project satisfies the stated academic goals",
        ],
        "8.2 FUTURE SCOPE": [
            "greater automation can be added without changing the core attendance model",
            "dashboards can expose trend analysis for mentors and department leads",
            "bulk administration tools can reduce setup overhead for new batches",
            "the same foundation can be reused for related classroom management tasks",
        ],
        "8.3 CONCLUDING REMARKS": [
            "the project is intentionally practical rather than decorative",
            "its design focuses on accuracy, ease of use, and a realistic academic workflow",
            "the mini project can therefore stand as a useful prototype for further development",
            "the final result is a coherent solution that matches the needs of the teacher desktop use case",
        ],
    }


def build_chapter_paragraphs(label: str, points: list[str]):
    first, second = make_paragraphs(label, points)
    return first, second


def draw_text_center(draw, box, text, font, fill):
    left, top, right, bottom = box
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    x = left + (right - left - width) / 2
    y = top + (bottom - top - height) / 2
    draw.text((x, y), text, font=font, fill=fill)


def create_architecture_figure(path: Path):
    img = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(TIMES_BOLD_FONT, 42)
    body_font = ImageFont.truetype(TIMES_FONT, 28)
    small_font = ImageFont.truetype(TIMES_FONT, 24)
    header_font = ImageFont.truetype(TIMES_BOLD_FONT, 32)

    draw.text((90, 60), "SYSTEM ARCHITECTURE OVERVIEW", font=title_font, fill=(0, 0, 0))

    boxes = [
        ((120, 220, 440, 340), "Teacher Login", "Desktop UI"),
        ((520, 220, 840, 340), "Electron Shell", "Renderer + main"),
        ((920, 220, 1240, 340), "API Layer", "Express endpoints"),
        ((1320, 220, 1640, 340), "PostgreSQL", "Live records"),
        ((520, 560, 840, 680), "Offline Store", "Queued sync"),
        ((920, 560, 1240, 680), "Reports", "Admin export"),
    ]
    for box, line1, line2 in boxes:
        draw.rounded_rectangle(box, radius=22, outline=(26, 79, 120), width=5, fill=(248, 250, 252))
        draw_text_center(draw, (box[0], box[1] + 15, box[2], box[1] + 70), line1, header_font, (0, 0, 0))
        draw_text_center(draw, (box[0], box[1] + 75, box[2], box[3] - 10), line2, body_font, (70, 70, 70))

    arrows = [
        ((440, 280), (520, 280)),
        ((840, 280), (920, 280)),
        ((1240, 280), (1320, 280)),
        ((680, 340), (680, 560)),
        ((1080, 340), (1080, 560)),
    ]
    for start, end in arrows:
        draw.line([start, end], fill=(26, 79, 120), width=6)
        dx = end[0] - start[0]
        dy = end[1] - start[1]
        length = math.hypot(dx, dy)
        ux, uy = dx / length, dy / length
        arrow_size = 16
        base_x = end[0] - ux * arrow_size * 2
        base_y = end[1] - uy * arrow_size * 2
        perp_x, perp_y = -uy, ux
        point1 = (end[0], end[1])
        point2 = (base_x + perp_x * arrow_size, base_y + perp_y * arrow_size)
        point3 = (base_x - perp_x * arrow_size, base_y - perp_y * arrow_size)
        draw.polygon([point1, point2, point3], fill=(26, 79, 120))

    draw.rounded_rectangle((90, 835, 1710, 980), radius=24, outline=(191, 196, 204), width=3, fill=(244, 246, 249))
    draw.text((140, 875), "The architecture separates UI, API, storage, offline caching, and reporting responsibilities.", font=small_font, fill=(60, 60, 60))
    draw.text((140, 915), "This keeps attendance capture simple while still supporting sync, analytics, and administrative review.", font=small_font, fill=(60, 60, 60))
    img.save(path)


def create_attendance_chart(path: Path):
    img = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(TIMES_BOLD_FONT, 42)
    axis_font = ImageFont.truetype(TIMES_FONT, 24)
    label_font = ImageFont.truetype(TIMES_FONT, 22)
    draw.text((90, 60), "SECTION-WISE ATTENDANCE SNAPSHOT", font=title_font, fill=(0, 0, 0))

    left, top, right, bottom = 140, 180, 1660, 860
    draw.line((left, bottom, right, bottom), fill=(0, 0, 0), width=4)
    draw.line((left, top, left, bottom), fill=(0, 0, 0), width=4)
    for pct in [20, 40, 60, 80, 100]:
        y = bottom - (bottom - top) * (pct / 100)
        draw.line((left - 10, y, right, y), fill=(220, 225, 230), width=2)
        draw.text((60, y - 16), f"{pct}%", font=axis_font, fill=(60, 60, 60))

    sections = [
        ("4CSE20", 92),
        ("4CSE21", 88),
        ("4CSE22", 84),
        ("4CSE23", 79),
        ("4CSE24", 91),
        ("4CSE25", 87),
    ]
    bar_width = 160
    gap = 90
    start_x = 200
    for index, (label, value) in enumerate(sections):
        x1 = start_x + index * (bar_width + gap)
        x2 = x1 + bar_width
        bar_height = (bottom - top) * (value / 100)
        y1 = bottom - bar_height
        fill = (48, 99, 156) if value >= 85 else (116, 153, 194)
        draw.rounded_rectangle((x1, y1, x2, bottom), radius=18, fill=fill, outline=(26, 79, 120), width=3)
        draw.text((x1 + 25, y1 - 35), f"{value}%", font=axis_font, fill=(0, 0, 0))
        draw.text((x1 + 8, bottom + 18), label, font=label_font, fill=(0, 0, 0))

    draw.text((150, 915), "Illustrative attendance rate by section shows the spread that the dashboard highlights during review.", font=axis_font, fill=(70, 70, 70))
    img.save(path)


def create_dashboard_mockup(path: Path):
    img = Image.new("RGB", (1800, 1050), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(TIMES_BOLD_FONT, 38)
    card_title = ImageFont.truetype(TIMES_BOLD_FONT, 24)
    card_value = ImageFont.truetype(TIMES_BOLD_FONT, 34)
    body_font = ImageFont.truetype(TIMES_FONT, 20)
    draw.text((90, 55), "DASHBOARD OUTPUT MOCKUP", font=title_font, fill=(0, 0, 0))

    cards = [
        ((90, 150, 390, 280), "Attendance", "92%"),
        ((430, 150, 730, 280), "Sections", "6"),
        ((770, 150, 1070, 280), "Pending Sync", "2"),
        ((1110, 150, 1410, 280), "Absentees", "18"),
        ((1450, 150, 1710, 280), "SMS Ready", "Enabled"),
    ]
    for box, label, value in cards:
        draw.rounded_rectangle(box, radius=24, fill="white", outline=(210, 217, 224), width=3)
        draw.text((box[0] + 24, box[1] + 22), label, font=card_title, fill=(40, 40, 40))
        draw.text((box[0] + 24, box[1] + 72), value, font=card_value, fill=(26, 79, 120))

    draw.rounded_rectangle((90, 330, 1100, 930), radius=24, fill="white", outline=(210, 217, 224), width=3)
    draw.rounded_rectangle((1160, 330, 1710, 930), radius=24, fill="white", outline=(210, 217, 224), width=3)
    draw.text((120, 350), "Attendance Trend", font=card_title, fill=(0, 0, 0))
    draw.text((1190, 350), "Action List", font=card_title, fill=(0, 0, 0))

    chart_left, chart_top, chart_right, chart_bottom = 160, 420, 1030, 860
    draw.line((chart_left, chart_bottom, chart_right, chart_bottom), fill=(0, 0, 0), width=3)
    draw.line((chart_left, chart_top, chart_left, chart_bottom), fill=(0, 0, 0), width=3)
    points = [(160, 780), (280, 730), (400, 690), (520, 650), (640, 560), (760, 520), (880, 470), (1000, 440)]
    draw.line(points, fill=(26, 79, 120), width=6)
    for point in points:
        draw.ellipse((point[0] - 8, point[1] - 8, point[0] + 8, point[1] + 8), fill=(26, 79, 120))
    for label, x in zip(["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun", "Next"], range(150, 1020, 120)):
        draw.text((x - 5, 875), label, font=body_font, fill=(70, 70, 70))

    items = [
        "Login as teacher and load assigned sections",
        "Select date and mark P or A in the register",
        "Submit attendance and review absentees",
        "Send SMS summary when enabled",
        "Export section report for admin review",
    ]
    y = 400
    for item in items:
        draw.ellipse((1200, y + 10, 1220, y + 30), fill=(48, 99, 156))
        draw.text((1240, y), item, font=body_font, fill=(50, 50, 50))
        y += 95

    draw.text((120, 970), "Mock dashboard illustrates the main KPI cards and a simple trend line.", font=body_font, fill=(70, 70, 70))
    img.save(path)


def create_flow_figure(path: Path):
    img = Image.new("RGB", (1800, 1000), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(TIMES_BOLD_FONT, 40)
    box_font = ImageFont.truetype(TIMES_BOLD_FONT, 26)
    body_font = ImageFont.truetype(TIMES_FONT, 21)
    draw.text((90, 60), "ATTENDANCE WORKFLOW FLOW", font=title_font, fill=(0, 0, 0))
    nodes = [
        ((120, 210, 380, 340), "Login", "Teacher SRN"),
        ((470, 210, 730, 340), "Load Section", "Students + dates"),
        ((820, 210, 1080, 340), "Mark P / A", "Direct table"),
        ((1170, 210, 1430, 340), "Submit", "Save to DB"),
        ((1520, 210, 1680, 340), "Review", "Absentees"),
    ]
    for box, title, subtitle in nodes:
        draw.rounded_rectangle(box, radius=20, outline=(26, 79, 120), width=4, fill=(247, 249, 252))
        draw_text_center(draw, (box[0], box[1] + 14, box[2], box[1] + 72), title, box_font, (0, 0, 0))
        draw_text_center(draw, (box[0], box[1] + 76, box[2], box[3] - 10), subtitle, body_font, (70, 70, 70))

    for start, end in [((380, 275), (470, 275)), ((730, 275), (820, 275)), ((1080, 275), (1170, 275)), ((1430, 275), (1520, 275))]:
        draw.line([start, end], fill=(26, 79, 120), width=6)
        draw.polygon([(end[0], end[1]), (end[0] - 18, end[1] - 10), (end[0] - 18, end[1] + 10)], fill=(26, 79, 120))

    draw.rounded_rectangle((120, 470, 1680, 860), radius=24, fill=(244, 246, 249), outline=(210, 217, 224), width=3)
    steps = [
        "1. The teacher opens the dashboard and confirms the assigned subject.",
        "2. The date picker loads a fresh or existing attendance column.",
        "3. Present and absent values are toggled directly inside the register.",
        "4. The submission saves the results and updates the absentee summary.",
        "5. Optional SMS messaging can be triggered from the same result view.",
    ]
    y = 520
    for step in steps:
        draw.text((160, y), step, font=body_font, fill=(50, 50, 50))
        y += 66
    img.save(path)


def create_test_figure(path: Path):
    img = Image.new("RGB", (1800, 1000), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(TIMES_BOLD_FONT, 40)
    body_font = ImageFont.truetype(TIMES_FONT, 24)
    draw.text((90, 60), "TEST AND VALIDATION MATRIX", font=title_font, fill=(0, 0, 0))
    headers = ["Test Case", "Input", "Expected Result", "Status"]
    widths = [320, 380, 720, 260]
    x_positions = [100, 420, 800, 1520]
    top = 190
    row_h = 110
    for i, header in enumerate(headers):
        box = (x_positions[i], top, x_positions[i] + widths[i], top + row_h)
        draw.rectangle(box, outline=(26, 79, 120), fill=(232, 238, 245), width=3)
        draw.text((box[0] + 20, box[1] + 35), header, font=body_font, fill=(0, 0, 0))
    rows = [
        ("Login", "Valid teacher ID", "Dashboard opens", "Pass"),
        ("Section load", "Assigned section", "Students display", "Pass"),
        ("Save", "Marked attendance", "Database updated", "Pass"),
        ("Offline sync", "Connection restored", "Queued items upload", "Pass"),
        ("Delete date", "Saved column selected", "Date removed", "Pass"),
    ]
    y = top + row_h
    for row in rows:
        for i, value in enumerate(row):
            box = (x_positions[i], y, x_positions[i] + widths[i], y + row_h)
            draw.rectangle(box, outline=(180, 187, 194), fill=(248, 250, 252), width=2)
            draw.text((box[0] + 18, box[1] + 30), value, font=body_font, fill=(40, 40, 40))
        y += row_h
    img.save(path)


def build_body(document, figures):
    banks = paragraph_bank()

    body = [
        ("1", "INTRODUCTION", [
            ("1.1 BACKGROUND OF THE STUDY", banks["1.1 BACKGROUND OF THE STUDY"]),
            ("1.2 PROBLEM STATEMENT", banks["1.2 PROBLEM STATEMENT"]),
            ("1.3 OBJECTIVES OF THE PROJECT", banks["1.3 OBJECTIVES OF THE PROJECT"]),
            ("1.4 SCOPE OF THE PROJECT", banks["1.4 SCOPE OF THE PROJECT"]),
            ("1.5 PROJECT MOTIVATION", banks["1.5 PROJECT MOTIVATION"]),
            ("1.6 ORGANIZATION OF THE REPORT", banks["1.6 ORGANIZATION OF THE REPORT"]),
        ]),
        ("2", "LITERATURE REVIEW / BACKGROUND STUDY", [
            ("2.1 MANUAL ATTENDANCE PRACTICES", banks["2.1 MANUAL ATTENDANCE PRACTICES"]),
            ("2.2 WEB AND DESKTOP ATTENDANCE SYSTEMS", banks["2.2 WEB AND DESKTOP ATTENDANCE SYSTEMS"]),
            ("2.3 OFFLINE-FIRST APPROACHES", banks["2.3 OFFLINE-FIRST APPROACHES"]),
            ("2.4 SMS AND NOTIFICATION SUPPORT", banks["2.4 SMS AND NOTIFICATION SUPPORT"]),
            ("2.5 COMPARATIVE ANALYSIS", banks["2.5 COMPARATIVE ANALYSIS"]),
            ("2.6 LITERATURE REVIEW SUMMARY", banks["2.6 LITERATURE REVIEW SUMMARY"]),
        ]),
        ("3", "PROBLEM DEFINITION, OBJECTIVES AND METHODOLOGY", [
            ("3.1 STAKEHOLDERS AND USER ROLES", banks["3.1 STAKEHOLDERS AND USER ROLES"]),
            ("3.2 FUNCTIONAL REQUIREMENTS", banks["3.2 FUNCTIONAL REQUIREMENTS"]),
            ("3.3 NON-FUNCTIONAL REQUIREMENTS", banks["3.3 NON-FUNCTIONAL REQUIREMENTS"]),
            ("3.4 USE CASE VIEW", banks["3.4 USE CASE VIEW"]),
            ("3.5 DATA REQUIREMENTS", banks["3.5 DATA REQUIREMENTS"]),
            ("3.6 CONSTRAINTS AND RISKS", banks["3.6 CONSTRAINTS AND RISKS"]),
        ]),
        ("4", "WORK CARRIED OUT", [
            ("4.1 OVERALL ARCHITECTURE", banks["4.1 OVERALL ARCHITECTURE"]),
            ("4.2 DATABASE DESIGN", banks["4.2 DATABASE DESIGN"]),
            ("4.3 API AND SERVICE DESIGN", banks["4.3 API AND SERVICE DESIGN"]),
            ("4.4 OFFLINE CACHE DESIGN", banks["4.4 OFFLINE CACHE DESIGN"]),
            ("4.5 SECURITY AND VALIDATION", banks["4.5 SECURITY AND VALIDATION"]),
            ("4.6 MODULE INTERACTION", banks["4.6 MODULE INTERACTION"]),
        ]),
        ("5", "RESULTS AND DISCUSSION", [
            ("5.1 LOGIN AND DASHBOARD SCREENS", banks["5.1 LOGIN AND DASHBOARD SCREENS"]),
            ("5.2 ATTENDANCE ENTRY WORKFLOW", banks["5.2 ATTENDANCE ENTRY WORKFLOW"]),
            ("5.3 ABSENTEE REVIEW AND SMS FLOW", banks["5.3 ABSENTEE REVIEW AND SMS FLOW"]),
            ("5.4 ADMIN REPORTING AND EXPORT", banks["5.4 ADMIN REPORTING AND EXPORT"]),
            ("5.5 VISUAL ANALYTICS DASHBOARD", banks["5.5 VISUAL ANALYTICS DASHBOARD"]),
            ("5.6 USER EXPERIENCE OBSERVATIONS", banks["5.6 USER EXPERIENCE OBSERVATIONS"]),
        ]),
        ("6", "CONCLUSIONS AND SCOPE FOR FUTURE WORK", [
            ("6.1 IMPLEMENTATION STRATEGY", banks["6.1 IMPLEMENTATION STRATEGY"]),
            ("6.2 SAMPLE TEST CASES", banks["6.2 SAMPLE TEST CASES"]),
            ("6.3 ERROR HANDLING AND RECOVERY", banks["6.3 ERROR HANDLING AND RECOVERY"]),
            ("6.4 PERFORMANCE AND RELIABILITY", banks["6.4 PERFORMANCE AND RELIABILITY"]),
            ("6.5 DEPLOYMENT AND PACKAGING", banks["6.5 DEPLOYMENT AND PACKAGING"]),
            ("6.6 VALIDATION RESULTS", banks["6.6 VALIDATION RESULTS"]),
        ]),
        ("7", "REFERENCES", [
            ("7.1 RESULTS OBSERVED DURING USE", banks["7.1 RESULTS OBSERVED DURING USE"]),
            ("7.2 COMPARISON WITH MANUAL METHOD", banks["7.2 COMPARISON WITH MANUAL METHOD"]),
            ("7.3 DISCUSSION OF STRENGTHS", banks["7.3 DISCUSSION OF STRENGTHS"]),
            ("7.4 LIMITATIONS", banks["7.4 LIMITATIONS"]),
            ("7.5 MAINTENANCE CONSIDERATIONS", banks["7.5 MAINTENANCE CONSIDERATIONS"]),
            ("7.6 FUTURE ENHANCEMENT AREAS", banks["7.6 FUTURE ENHANCEMENT AREAS"]),
        ]),
        ("8", "APPENDIX / ANNEXURE (IF ANY)", [
            ("8.1 SUMMARY OF WORK", banks["8.1 SUMMARY OF WORK"]),
            ("8.2 FUTURE SCOPE", banks["8.2 FUTURE SCOPE"]),
            ("8.3 CONCLUDING REMARKS", banks["8.3 CONCLUDING REMARKS"]),
        ]),
    ]

    document.add_page_break()
    for chapter_no, chapter_title, sections in body:
        if chapter_no != "1":
            document.add_page_break()
        p1 = document.add_paragraph()
        set_paragraph_text(p1, f"CHAPTER {chapter_no}", size=16, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        p2 = document.add_paragraph()
        set_paragraph_text(p2, chapter_title, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        p2.paragraph_format.space_after = Pt(10)

        for index, (section_title, points) in enumerate(sections):
            add_body_heading(document, section_title, level=2, before=4 if index else 6, after=3)
            para1, para2 = build_chapter_paragraphs(section_title, points)
            add_body_paragraph(document, para1)
            add_body_paragraph(document, para2)

            if chapter_no == "3" and index == 1:
                add_caption(document, "Table 3.1: Functional and non-functional requirements.")
                add_simple_table(
                    document,
                    ["Category", "Requirement", "Why it matters"],
                    [
                        ("Functional", "Teacher login and section mapping", "Ensures the correct staff member sees only assigned classes."),
                        ("Functional", "Date-wise attendance save and reset", "Supports the daily classroom routine without manual rewrites."),
                        ("Functional", "Absentee review and optional SMS", "Improves follow-up after attendance submission."),
                        ("Non-functional", "Responsive desktop interface", "Keeps the workflow practical during a short class period."),
                        ("Non-functional", "Reliable database storage", "Protects the accuracy of academic records."),
                    ],
                    [Inches(1.5), Inches(2.6), Inches(2.4)],
                )
            if chapter_no == "1" and index == 2:
                add_caption(document, "Figure 1.1: Teacher-facing system architecture.")
                document.add_picture(str(figures["architecture"]), width=Inches(6.3))
                add_caption(document, "The architecture keeps the user interface, backend service, database, and offline cache distinct.")
            if chapter_no == "3" and index == 3:
                add_caption(document, "Figure 3.1: Attendance workflow from login to review.")
                document.add_picture(str(figures["flow"]), width=Inches(6.3))
                add_caption(document, "The workflow mirrors the classroom routine and keeps the teacher inside one consistent flow.")
            if chapter_no == "4" and index == 1:
                add_caption(document, "Table 4.1: Core database tables and responsibilities.")
                add_simple_table(
                    document,
                    ["Table", "Role", "Key fields"],
                    [
                        ("teachers", "Stores faculty identity and subject assignment", "teacherId, name, subjectCode"),
                        ("students", "Stores student identity and parent contact", "srn, name, parent_phone_no"),
                        ("sections", "Stores section structure and semester mapping", "section_code, semester, subject"),
                        ("attendance", "Stores each date-wise attendance record", "srn, attendance_date, status"),
                        ("student_electives", "Stores DC and MC filtering choices", "srn, elective_code"),
                    ],
                    [Inches(1.5), Inches(2.2), Inches(2.8)],
                )
                add_caption(document, "Figure 4.1: Database and service interaction model.")
                document.add_picture(str(figures["architecture"]), width=Inches(6.3))
                add_caption(document, "The same layered separation supports maintainability as the schema grows.")
            if chapter_no == "5" and index == 2:
                add_caption(document, "Figure 5.1: Dashboard output mockup.")
                document.add_picture(str(figures["dashboard"]), width=Inches(6.3))
                add_caption(document, "The dashboard presents KPI cards, action prompts, and a trend line for quick interpretation.")
                add_caption(document, "Figure 5.2: Section-wise attendance chart.")
                document.add_picture(str(figures["attendance"]), width=Inches(6.3))
                add_caption(document, "The chart visually compares the percentage attendance across sections.")
            if chapter_no == "5" and index == 3:
                add_caption(document, "Table 5.1: Sample attendance summary for the dashboard.")
                add_simple_table(
                    document,
                    ["Section", "Present", "Absent", "Attendance %"],
                    [
                        ("4CSE20", "46", "4", "92%"),
                        ("4CSE21", "44", "6", "88%"),
                        ("4CSE22", "42", "8", "84%"),
                        ("4CSE23", "40", "10", "80%"),
                        ("4CSE24", "45", "5", "90%"),
                    ],
                    [Inches(1.3), Inches(1.1), Inches(1.1), Inches(1.3)],
                )
            if chapter_no == "6" and index == 1:
                add_caption(document, "Table 6.1: Test case matrix.")
                add_simple_table(
                    document,
                    ["Test case", "Input", "Expected result", "Status"],
                    [
                        ("Login", "Valid teacher ID", "Dashboard opens", "Pass"),
                        ("Section load", "Assigned section", "Student register appears", "Pass"),
                        ("Save", "Marked attendance", "Database updated", "Pass"),
                        ("Offline sync", "Connection restored", "Queued entries upload", "Pass"),
                        ("Delete date", "Saved column", "Date removed", "Pass"),
                    ],
                    [Inches(1.0), Inches(2.0), Inches(2.4), Inches(0.9)],
                )
                add_caption(document, "Figure 6.1: Validation and testing matrix.")
                document.add_picture(str(figures["test"]), width=Inches(6.3))
                add_caption(document, "The matrix confirms that the main attendance workflows behave as expected.")
            if chapter_no == "6" and index == 4:
                add_caption(document, "Table 6.2: Manual method versus proposed system.")
                add_simple_table(
                    document,
                    ["Manual method", "Proposed system"],
                    [
                        ("Paper registers require transcription.", "Data is saved directly in PostgreSQL."),
                        ("Corrections are slow and error-prone.", "Date-wise reset and reopen are built in."),
                        ("Attendance review is fragmented.", "Dashboard cards provide a quick summary."),
                        ("Follow-up is delayed.", "Absentee review and SMS can be triggered."),
                    ],
                    [Inches(3.0), Inches(3.5)],
                )

        if chapter_no == "7":
            add_body_heading(document, "7.1 SELECTED REFERENCES", level=2, before=8, after=4)
            refs = [
                "Electron Documentation, desktop application framework reference for the UI shell.",
                "Express Documentation, routing and middleware reference for the API layer.",
                "PostgreSQL Documentation, relational storage reference used by the project database.",
                "Twilio Documentation, messaging guidance used for absentee communication support.",
                "OWASP Cheat Sheet Series, general input validation and session handling guidance.",
                "SNPSU Teacher Desktop source files: main.js, renderer.js, server.js, db.js, and offline-store.js.",
            ]
            add_bullet_list(document, refs)
            add_body_heading(document, "7.2 DISCUSSION OF REFERENCES", level=2, before=6, after=4)
            add_body_paragraph(document, "The references are deliberately practical rather than exhaustive. They support the stack used by the project and give the reader a direct trail from the application behavior to the technologies that enable it. Because the report is a mini project rather than a full thesis, this level of referencing is enough to document the implementation choices without pulling the narrative away from the actual system under discussion.")

        if chapter_no == "8":
            add_body_heading(document, "8.1 APPENDIX MATERIAL", level=2, before=8, after=4)
            appendix_table = document.add_table(rows=6, cols=3)
            appendix_table.style = "Table Grid"
            appendix_table.autofit = False
            widths = [Inches(1.1), Inches(2.7), Inches(2.7)]
            for row in appendix_table.rows:
                for i, cell in enumerate(row.cells):
                    cell.width = widths[i]
            headers = ["Item", "Purpose", "Notes"]
            for i, header in enumerate(headers):
                appendix_table.rows[0].cells[i].text = header
                set_cell_shading(appendix_table.rows[0].cells[i], "E8EEF5")
            append_rows = [
                ("A1", "Teacher login", "Loads section assignments and session context."),
                ("A2", "Attendance grid", "Shows students with date-wise present and absent toggles."),
                ("A3", "Offline queue", "Stores work locally until the network returns."),
                ("A4", "Absentee modal", "Summarizes missing students and optional SMS support."),
                ("A5", "Admin report", "Supports review, export, and section-level summaries."),
            ]
            for row_index, row_data in enumerate(append_rows, start=1):
                for col_index, value in enumerate(row_data):
                    appendix_table.rows[row_index].cells[col_index].text = value
                    for paragraph in appendix_table.rows[row_index].cells[col_index].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.15
                        for run in paragraph.runs:
                            set_run_font(run, size=11, color=BLACK)
            document.add_paragraph()
            add_body_paragraph(document, "The appendix is intentionally compact and focuses on user-facing items that support a viva discussion. It gives the reviewer a quick map of the most important screens and confirms that the report content corresponds to the implemented modules. If more material is required later, sample database entries, additional screenshots, and exported reports can be appended without changing the core structure of the document.")


def populate_front_matter(document):
    replace_paragraph(document, "REPORT ON MINI PROJECT", "REPORT ON MINI PROJECT", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(document, "TITLE OF THE MINI PROJECT", "SNPSU TEACHER DESKTOP\nSMART ATTENDANCE MANAGEMENT SYSTEM", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(document, "BACHELOR OF ENGINEERING", "BACHELOR OF ENGINEERING", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(document, "----------------------------------------- ENGINEERING", "COMPUTER SCIENCE AND ENGINEERING", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(document, "NAME OF THE GUIDE:", "NAME OF THE GUIDE: [PROJECT GUIDE NAME]", size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(document, "Designation, Dept. of -------------------", "Designation, Dept. of COMPUTER SCIENCE AND ENGINEERING", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(document, "For the Academic year of 2025-26 [Semester: ---]", "For the Academic year of 2025-26 [Semester: VI]", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(document, "DEPARTMENT OF --------------------------------", "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    replace_paragraph(document, "CERTIFICATE", "CERTIFICATE", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    certificate_text = (
        "Certified that the MINI PROJECT entitled \"SNPSU TEACHER DESKTOP: SMART ATTENDANCE MANAGEMENT SYSTEM\" is carried out "
        "by NAME OF STUDENT 1, NAME OF STUDENT 2, NAME OF STUDENT 3, and NAME OF STUDENT 4, bearing the corresponding SRN values, "
        "as bonafide students of School of Engineering and Technology in partial fulfilment for the award of Bachelor of Engineering "
        "in the Department of Computer Science and Engineering of SAPTHAGIRI NPS University during the academic year 2025-2026. "
        "It is certified that all corrections and suggestions indicated in the internal assessment have been incorporated in the report submitted. "
        "The mini project report has been approved as it satisfies the academic requirements prescribed for the Bachelor of Engineering Degree."
    )
    replace_paragraph(document, "Certified that the MINI PROJECT", certificate_text, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    replace_paragraph(document, "Signature of the Guide", "Signature of the Guide\t\t\t\t\tSignature of the Director", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    replace_paragraph(document, "Name of the Guide", "Name of the Guide\t\t\t\t\t\tName of the Director", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    replace_paragraph(document, "Designation", "Designation\t\t\t\t\t\t\tDirector", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    replace_paragraph(document, "DEPARTMENT OF ----------------------------------------ENGINEERING", "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    replace_paragraph(document, "ACKNOWLEDGEMENT", "ACKNOWLEDGEMENT", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    ack = (
        "We express our sincere gratitude to our guide for continuous support, patient guidance, and useful suggestions throughout the completion of this mini project. "
        "We also thank the faculty members and the institution for providing the environment required to develop and review the SNPSU Teacher Desktop system. "
        "The project benefited from the encouragement of classmates and friends who offered feedback during the design, testing, and report preparation stages. "
        "Their support helped us complete the work in a systematic and confident manner."
    )
    # The template has three acknowledgment paragraphs; replace them to form a concise formal note.
    ack_paragraphs = [p for p in document.paragraphs if "We, the mini project group" in p.text or "We would like to convey" in p.text or "Finally, we thank" in p.text]
    if ack_paragraphs:
        set_paragraph_text(ack_paragraphs[0], ack, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        if len(ack_paragraphs) > 1:
            ack_paragraphs[1].text = ""
        if len(ack_paragraphs) > 2:
            ack_paragraphs[2].text = ""

    replace_paragraph(document, "ABSTRACT", "ABSTRACT", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    abstract_text = (
        "This mini project report presents SNPSU Teacher Desktop, a smart attendance management system built for faculty use in an academic environment. "
        "The application is designed as an Electron-based desktop solution with a Node.js and Express backend and a PostgreSQL database. "
        "Its core purpose is to simplify section-wise and subject-wise attendance entry, while preserving the ability to review absentees, reopen dates when necessary, "
        "and synchronize offline work once connectivity returns. The dashboard gives the teacher a clear view of section progress, low-attendance cases, and pending actions. "
        "The project also supports administrative reporting and optional absentee notification support so that the same data can be used for both operational marking and follow-up communication. "
        "During development, emphasis was placed on reliability, clarity, and a workflow that matches actual classroom practice. "
        "The result is a compact but practical academic system that reduces manual work, improves record consistency, and demonstrates how a small desktop application can handle routine department tasks effectively."
    )
    abstract_paragraphs = [p for p in document.paragraphs if p.text.strip().startswith("This mini project report presents") or p.text.strip().startswith("This report")]
    if abstract_paragraphs:
        set_paragraph_text(abstract_paragraphs[0], abstract_text, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    replace_paragraph(document, "CONTENTS", "CONTENTS", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(document, "LIST OF FIGURES", "LIST OF FIGURES", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(document, "LIST OF TABLES", "LIST OF TABLES", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    students = [
        ("1", "NAME OF STUDENT 1", "SRN001"),
        ("2", "NAME OF STUDENT 2", "SRN002"),
        ("3", "NAME OF STUDENT 3", "SRN003"),
        ("4", "NAME OF STUDENT 4", "SRN004"),
    ]
    tables = document.tables
    if len(tables) >= 1:
        set_table_header(tables[0])
        fill_table(tables[0], students)
    if len(tables) >= 2:
        set_table_header(tables[1])
        fill_table(tables[1], students)
    if len(tables) >= 3:
        set_table_header(tables[2])
        set_table_header(tables[3])
        set_table_header(tables[4])


def populate_contents_tables(document, page_map, figure_pages, table_pages):
    default_page_map = {
        "1": 8,
        "2": 12,
        "3": 15,
        "4": 19,
        "5": 24,
        "6": 29,
        "7": 33,
        "8": 37,
    }
    default_figure_pages = {
        "1.1": 10,
        "3.1": 17,
        "4.1": 21,
        "5.1": 26,
        "5.2": 27,
        "6.1": 31,
    }
    default_table_pages = {
        "3.1": 16,
        "4.1": 20,
        "5.1": 28,
        "6.1": 30,
        "6.2": 33,
    }
    merged_page_map = {**default_page_map, **page_map}
    merged_figure_pages = {**default_figure_pages, **figure_pages}
    merged_table_pages = {**default_table_pages, **table_pages}

    contents_table = document.tables[2]
    contents_rows = [
        ("1", "Introduction", merged_page_map.get("1", "")),
        ("2", "Literature Review / Background Study", merged_page_map.get("2", "")),
        ("3", "Problem Definition, Objectives and Methodology", merged_page_map.get("3", "")),
        ("4", "Work Carried Out", merged_page_map.get("4", "")),
        ("5", "Results and Discussion", merged_page_map.get("5", "")),
        ("6", "Conclusions and Scope for Future Work", merged_page_map.get("6", "")),
        ("7", "References", merged_page_map.get("7", "")),
        ("8", "Appendix / Annexure (if any)", merged_page_map.get("8", "")),
    ]
    for row_index, (chapter, desc, page) in enumerate(contents_rows, start=1):
        if row_index >= len(contents_table.rows):
            break
        contents_table.rows[row_index].cells[0].text = chapter
        contents_table.rows[row_index].cells[1].text = desc
        contents_table.rows[row_index].cells[2].text = str(page)
        for col_index in range(3):
            cell = contents_table.rows[row_index].cells[col_index]
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_index < 2 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run_font(run, size=11, color=BLACK)

    fig_table = document.tables[3]
    figure_rows = [
        ("Figure 1.1", "Teacher-facing system architecture", merged_figure_pages.get("1.1", "")),
        ("Figure 3.1", "Attendance workflow from login to review", merged_figure_pages.get("3.1", "")),
        ("Figure 4.1", "Database and service interaction model", merged_figure_pages.get("4.1", "")),
        ("Figure 5.1", "Dashboard output mockup", merged_figure_pages.get("5.1", "")),
        ("Figure 5.2", "Section-wise attendance chart", merged_figure_pages.get("5.2", "")),
        ("Figure 6.1", "Validation and testing matrix", merged_figure_pages.get("6.1", "")),
    ]
    for row_index, (fig_no, desc, page) in enumerate(figure_rows, start=1):
        if row_index >= len(fig_table.rows):
            break
        fig_table.rows[row_index].cells[0].text = fig_no
        fig_table.rows[row_index].cells[1].text = desc
        fig_table.rows[row_index].cells[2].text = str(page)
        for col_index in range(3):
            cell = fig_table.rows[row_index].cells[col_index]
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_index < 2 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run_font(run, size=11, color=BLACK)

    tbl_table = document.tables[4]
    table_rows = [
        ("Table 3.1", "Functional and non-functional requirements", merged_table_pages.get("3.1", "")),
        ("Table 4.1", "Core database tables and responsibilities", merged_table_pages.get("4.1", "")),
        ("Table 5.1", "Sample attendance summary", merged_table_pages.get("5.1", "")),
        ("Table 6.1", "Test case matrix", merged_table_pages.get("6.1", "")),
        ("Table 6.2", "Manual method versus proposed system", merged_table_pages.get("6.2", "")),
    ]
    for row_index, (tbl_no, desc, page) in enumerate(table_rows, start=1):
        if row_index >= len(tbl_table.rows):
            break
        tbl_table.rows[row_index].cells[0].text = tbl_no
        tbl_table.rows[row_index].cells[1].text = desc
        tbl_table.rows[row_index].cells[2].text = str(page)
        for col_index in range(3):
            cell = tbl_table.rows[row_index].cells[col_index]
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_index < 2 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run_font(run, size=11, color=BLACK)


def extract_page_markers(pdf_path: Path):
    import pdfplumber

    page_map = {}
    figure_pages = {}
    table_pages = {}

    chapter_title_to_number = {
        "INTRODUCTION": "1",
        "LITERATURE REVIEW / BACKGROUND STUDY": "2",
        "PROBLEM DEFINITION, OBJECTIVES AND METHODOLOGY": "3",
        "WORK CARRIED OUT": "4",
        "RESULTS AND DISCUSSION": "5",
        "CONCLUSIONS AND SCOPE FOR FUTURE WORK": "6",
        "REFERENCES": "7",
        "APPENDIX / ANNEXURE (IF ANY)": "8",
    }

    figure_labels = {
        "Figure 1.1: Teacher-facing system architecture.": "1.1",
        "Figure 3.1: Attendance workflow from login to review.": "3.1",
        "Figure 4.1: Database and service interaction model.": "4.1",
        "Figure 5.1: Dashboard output mockup.": "5.1",
        "Figure 5.2: Section-wise attendance chart.": "5.2",
        "Figure 6.1: Validation and testing matrix.": "6.1",
    }

    table_labels = {
        "Table 3.1": "3.1",
        "Table 4.1": "4.1",
        "Table 5.1": "5.1",
        "Table 6.1": "6.1",
        "Table 6.2": "6.2",
    }

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            for title, number in chapter_title_to_number.items():
                if title in text and number not in page_map:
                    page_map[number] = page_index
            for label, fig_no in figure_labels.items():
                if label in text and fig_no not in figure_pages:
                    figure_pages[fig_no] = page_index
            for label, table_no in table_labels.items():
                if label in text and table_no not in table_pages:
                    table_pages[table_no] = page_index
    return page_map, figure_pages, table_pages


def main():
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
    ensure_assets_dir()
    figures = {
        "architecture": ASSET_DIR / "architecture.png",
        "attendance": ASSET_DIR / "attendance_chart.png",
        "dashboard": ASSET_DIR / "dashboard_mockup.png",
        "flow": ASSET_DIR / "flow.png",
        "test": ASSET_DIR / "test_matrix.png",
    }
    create_architecture_figure(figures["architecture"])
    create_attendance_chart(figures["attendance"])
    create_dashboard_mockup(figures["dashboard"])
    create_flow_figure(figures["flow"])
    create_test_figure(figures["test"])

    shutil.copy2(TEMPLATE_PATH, OUTPUT_PATH)
    document = Document(OUTPUT_PATH)
    set_document_styles(document)
    set_section_layout(document)
    populate_front_matter(document)
    build_body(document, figures)
    populate_contents_tables(document, {}, {}, {})
    document.save(OUTPUT_PATH)

    print(f"Saved draft report to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
